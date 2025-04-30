from django.shortcuts import render
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.core.files.storage import FileSystemStorage
import os
from docx2pdf import convert
import pythoncom
from pdf2docx import Converter
from docx import Document
from django.core.files.storage import default_storage
from fpdf import FPDF
import PyPDF2
import uuid
from pdf2image import convert_from_path
from PIL import Image
from PyPDF2 import PdfMerger
from win32com.client import Dispatch
from xlsx2html import xlsx2html
import pdfkit

# Main File
def main(request):
    return render(request, 'main.html')

# Convert DOCX to PDF
@csrf_exempt
def convert_to_pdf(request):
    if request.method == 'POST' and request.FILES.get('file'):
        file = request.FILES['file']

        fs = FileSystemStorage(location='media/uploads')
        filename = fs.save(file.name, file)
        uploaded_file_path = fs.path(filename)

        output_dir = os.path.join('media', 'converted')
        os.makedirs(output_dir, exist_ok=True)

        base_filename = os.path.splitext(os.path.basename(file.name))[0]
        converted_file_path = os.path.join(output_dir, base_filename + '.pdf')

        try:
            pythoncom.CoInitialize()

            convert(uploaded_file_path, converted_file_path)

            pythoncom.CoUninitialize()

            converted_file_url = os.path.join('/media/converted/', os.path.basename(converted_file_path))
            return JsonResponse({
                'success': True,
                'download_link': converted_file_url
            })
        except Exception as e:
            return JsonResponse({'success': False, 'message': f'Error converting file: {str(e)}'})

    return JsonResponse({'success': False, 'message': 'File upload failed.'})

# Convert PDF to DOCX
@csrf_exempt
def convert_to_docx(request):
    if request.method == 'POST' and request.FILES.get('file'):
        file = request.FILES['file']

        fs = FileSystemStorage(location='media/uploads')
        filename = fs.save(file.name, file)
        uploaded_file_path = fs.path(filename)

        output_dir = os.path.join('media', 'converted')
        os.makedirs(output_dir, exist_ok=True)

        base_filename = os.path.splitext(os.path.basename(file.name))[0]
        converted_file_path = os.path.join(output_dir, base_filename + '.docx')

        try:
            cv = Converter(uploaded_file_path)
            cv.convert(converted_file_path, start=0, end=None)
            cv.close()

            converted_file_url = os.path.join('/media/converted/', os.path.basename(converted_file_path))
            return JsonResponse({
                'success': True,
                'download_link': converted_file_url
            })
        except Exception as e:
            return JsonResponse({'success': False, 'message': f'Error converting file: {str(e)}'})

    return JsonResponse({'success': False, 'message': 'File upload failed.'})

# Convert DOCX to TXT 
@csrf_exempt
def convert_to_txt(request):
    if request.method == 'POST' and request.FILES.get('file'):
        uploaded_file = request.FILES['file']

        file_path = default_storage.save('temp/' + uploaded_file.name, uploaded_file)

        doc = Document(default_storage.path(file_path))
        text_content = '\n'.join([para.text for para in doc.paragraphs])

        from django.core.files.base import ContentFile
        txt_filename = uploaded_file.name.replace('.docx', '.txt')
        txt_content = ContentFile(text_content)
        txt_path = default_storage.save('converted/' + txt_filename, txt_content)

        download_url = '/media/' + txt_path

        default_storage.delete(file_path)

        return JsonResponse({'download_url': download_url})

    return JsonResponse({'error': 'Invalid request'}, status=400)

# TXT to PDF
@csrf_exempt
def convert_txt_to_pdf(request):
    if request.method == 'POST' and request.FILES.get('file'):
        uploaded_file = request.FILES['file']

        fs = FileSystemStorage(location='media/uploads')
        filename = fs.save(uploaded_file.name, uploaded_file)
        uploaded_file_path = fs.path(filename)

        output_dir = os.path.join('media', 'converted')
        os.makedirs(output_dir, exist_ok=True)

        base_filename = os.path.splitext(os.path.basename(uploaded_file.name))[0]
        output_pdf_path = os.path.join(output_dir, base_filename + '.pdf')

        try:
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_font("Arial", size=12)

            with open(uploaded_file_path, 'r', encoding='utf-8') as file:
                for line in file:
                    pdf.multi_cell(0, 10, line.strip())

            pdf.output(output_pdf_path)

            converted_file_url = os.path.join('/media/converted/', os.path.basename(output_pdf_path))
            return JsonResponse({
                'success': True,
                'download_link': converted_file_url
            })
        except Exception as e:
            return JsonResponse({'success': False, 'message': f'Error converting TXT to PDF: {str(e)}'})

    return JsonResponse({'success': False, 'message': 'File upload failed.'})

# Convert TXT to DOCX
@csrf_exempt
def convert_txt_to_docx(request):
    if request.method == 'POST' and request.FILES.get('file'):
        file = request.FILES['file']

        fs = FileSystemStorage(location='media/uploads')
        filename = fs.save(file.name, file)
        uploaded_file_path = fs.path(filename)

        output_dir = os.path.join('media', 'converted')
        os.makedirs(output_dir, exist_ok=True)

        base_filename = os.path.splitext(os.path.basename(file.name))[0]
        converted_file_path = os.path.join(output_dir, base_filename + '.docx')

        try:
            with open(uploaded_file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()

            document = Document()

            for line in text_content.splitlines():
                document.add_paragraph(line)
            document.save(converted_file_path)

            converted_file_url = os.path.join('/media/converted/', os.path.basename(converted_file_path))
            return JsonResponse({
                'success': True,
                'download_link': converted_file_url
            })
        except Exception as e:
            return JsonResponse({'success': False, 'message': f'Error converting file: {str(e)}'})

    return JsonResponse({'success': False, 'message': 'File upload failed.'})

# Convert PDF to TXT
@csrf_exempt
def convert_pdf_to_txt(request):
    if request.method == 'POST' and request.FILES.get('file'):
        file = request.FILES['file']

        fs = FileSystemStorage(location='media/uploads')
        filename = fs.save(file.name, file)
        uploaded_file_path = fs.path(filename)

        output_dir = os.path.join('media', 'converted')
        os.makedirs(output_dir, exist_ok=True)

        base_filename = os.path.splitext(os.path.basename(file.name))[0]
        converted_file_path = os.path.join(output_dir, base_filename + '.txt')

        try:
            with open(uploaded_file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                text = ''
                for page in pdf_reader.pages:
                    text += page.extract_text() or ''

            with open(converted_file_path, 'w', encoding='utf-8') as txt_file:
                txt_file.write(text)

            converted_file_url = os.path.join('/media/converted/', os.path.basename(converted_file_path))
            return JsonResponse({
                'success': True,
                'download_link': converted_file_url
            })
        except Exception as e:
            return JsonResponse({'success': False, 'message': f'Error converting file: {str(e)}'})

    return JsonResponse({'success': False, 'message': 'File upload failed.'})

# Convert PDF to Images
@csrf_exempt
def convert_pdf_to_images(request):
    if request.method == 'POST' and request.FILES.get('file'):
        file = request.FILES['file']
        
        fs = FileSystemStorage(location='media/uploads')
        filename = fs.save(file.name, file)
        uploaded_file_path = fs.path(filename)
        
        output_dir = os.path.join('media', 'converted_images')
        os.makedirs(output_dir, exist_ok=True)

        try:
            poppler_path = r'C:\Users\User\OneDrive\Documents\ALEC\Python File Converter\my_project\app\Release-24.08.0-0\poppler-24.08.0\Library\bin'
            print("Using Poppler path:", poppler_path)

            pages = convert_from_path(
                uploaded_file_path,
                dpi=200,
                poppler_path=poppler_path
            )

            image_paths_png = []
            image_paths_jpg = []
            for idx, page in enumerate(pages):
                # PNG File
                image_filename_png = f"{uuid.uuid4()}.png"
                image_path_png = os.path.join(output_dir, image_filename_png)
                page.save(image_path_png, 'PNG')
                image_paths_png.append('/media/converted_images/' + image_filename_png)

                # JPG File
                image_filename_jpg = f"{uuid.uuid4()}.jpg"
                image_path_jpg = os.path.join(output_dir, image_filename_jpg)
                page.save(image_path_jpg, 'JPEG')
                image_paths_jpg.append('/media/converted_images/' + image_filename_jpg)

            return JsonResponse({
                'success': True,
                'images': {
                    'png': image_paths_png,
                    'jpg': image_paths_jpg
                }
            })

        except Exception as e:
            return JsonResponse({'success': False, 'message': f"Error: {str(e)}"})

    return JsonResponse({'success': False, 'message': 'No file uploaded'})

# Convert JPG/PNG to PDF
@csrf_exempt
def upload_images_to_pdf(request):
    if request.method == 'POST' and request.FILES.getlist('files'):
        files = request.FILES.getlist('files')

        upload_dir = 'media/uploads'
        output_dir = 'media/converted'
        os.makedirs(upload_dir, exist_ok=True)
        os.makedirs(output_dir, exist_ok=True)

        fs = FileSystemStorage(location=upload_dir)
        temp_pdf_paths = []

        for f in files:
            filename = fs.save(f.name, f)
            uploaded_file_path = fs.path(filename)

            ext = os.path.splitext(f.name)[1].lower()
            if ext in ['.jpg', '.jpeg', '.png']:
                try:
                    img = Image.open(uploaded_file_path)
                    img = img.convert('RGB')
                    temp_pdf_path = uploaded_file_path + '.pdf'
                    img.save(temp_pdf_path, format='PDF')
                    temp_pdf_paths.append(temp_pdf_path)
                except Exception as e:
                    return JsonResponse({'success': False, 'message': f'Error processing image: {str(e)}'})
            elif ext == '.pdf':
                temp_pdf_paths.append(uploaded_file_path)
            else:
                return JsonResponse({'success': False, 'message': f'Unsupported file type: {f.name}'})

        if not temp_pdf_paths:
            return JsonResponse({'success': False, 'message': 'No valid files to convert.'})

        base_filename = os.path.splitext(os.path.basename(files[0].name))[0]
        converted_filename = base_filename + '_merged.pdf'
        converted_file_path = os.path.join(output_dir, converted_filename)

        try:
            merger = PdfMerger()
            for pdf in temp_pdf_paths:
                merger.append(pdf)
            merger.write(converted_file_path)
            merger.close()

            converted_file_url = f'/media/converted/{converted_filename}'
            return JsonResponse({
                'success': True,
                'download_link': converted_file_url
            })
        except Exception as e:
            return JsonResponse({'success': False, 'message': f'Error merging PDFs: {str(e)}'})

    return JsonResponse({'success': False, 'message': 'No files uploaded.'})

# Convert PPTX to PDF
@csrf_exempt
def convert_pptx_to_pdf(request):
    if request.method == 'POST' and request.FILES.get('file'):
        file = request.FILES['file']

        fs = FileSystemStorage(location='media/uploads')
        filename = fs.save(file.name, file)
        uploaded_file_path = fs.path(filename)

        output_dir = os.path.join('media', 'converted')
        os.makedirs(output_dir, exist_ok=True)

        base_filename = os.path.splitext(os.path.basename(file.name))[0]
        converted_file_path = os.path.join(output_dir, base_filename + '.pdf')

        uploaded_file_path = os.path.abspath(uploaded_file_path).replace('/', '\\')
        converted_file_path = os.path.abspath(converted_file_path).replace('/', '\\')

        try:
            pythoncom.CoInitialize()

            powerpoint = Dispatch('PowerPoint.Application')
            powerpoint.Visible = 1

            presentation = powerpoint.Presentations.Open(uploaded_file_path, WithWindow=False)
            presentation.SaveAs(converted_file_path, 32)
            presentation.Close()
            powerpoint.Quit()

            pythoncom.CoUninitialize()

            converted_file_url = os.path.join('/media/converted/', os.path.basename(converted_file_path))
            return JsonResponse({
                'success': True,
                'download_link': converted_file_url
            })

        except Exception as e:
            pythoncom.CoUninitialize()
            return JsonResponse({'success': False, 'message': f'Error converting PPTX: {str(e)}'})

    return JsonResponse({'success': False, 'message': 'File upload failed.'})

# Convert XLSX to PDF
@csrf_exempt
def convert_xlsx_to_pdf(request):
    if request.method == 'POST' and request.FILES.get('file'):
        file = request.FILES['file']
        
        fs = FileSystemStorage(location='media/uploads')
        filename = fs.save(file.name, file)
        uploaded_file_path = fs.path(filename)

        output_dir = os.path.join('media', 'converted')
        os.makedirs(output_dir, exist_ok=True)

        base_filename = os.path.splitext(os.path.basename(file.name))[0]
        html_temp_path = os.path.join(output_dir, base_filename + '.html')
        pdf_output_path = os.path.join(output_dir, base_filename + '.pdf')

        try:
            with open(html_temp_path, "w", encoding='utf-8') as f:
                xlsx2html(uploaded_file_path, f)

            config = pdfkit.configuration(wkhtmltopdf=r'C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe')

            pdfkit.from_file(html_temp_path, pdf_output_path, configuration=config)

            os.remove(html_temp_path)
            os.remove(uploaded_file_path)

            download_link = os.path.join('/media/converted/', os.path.basename(pdf_output_path))

            return JsonResponse({
                'success': True,
                'download_link': download_link
            })

        except Exception as e:
            return JsonResponse({'success': False, 'message': f'Error converting file: {str(e)}'})

    return JsonResponse({'success': False, 'message': 'File upload failed.'})